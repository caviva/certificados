from flask import Flask, render_template, request, send_file
import pandas as pd
from docx import Document
import os
import tempfile
import subprocess

app = Flask(__name__)
df_global = None
UPLOAD_FOLDER = "uploads"
TEMPLATE_PATH = "plantilla/base.docx"

# Función segura para formatear montos
def formato_moneda(valor):
    try:
        return f"${float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return str(valor)

@app.route("/", methods=["GET", "POST"])
def index():
    global df_global
    if request.method == "POST":
        file = request.files["file"]
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        df_global = pd.read_excel(filepath)
        records = df_global.to_dict(orient="records")
        return render_template("index.html", data=records)
    return render_template("index.html", data=None)

@app.route("/generar/<int:row_id>")
def generar(row_id):
    global df_global
    if df_global is None or row_id >= len(df_global):
        return "Registro no encontrado", 404

    # Imprimir las columnas disponibles para depuración
    print("Columnas disponibles:", df_global.columns.tolist())
    
    row = df_global.iloc[row_id].fillna("")
    
    # Imprimir los valores específicos que no están funcionando
    print(f"fecha_cdp: '{row.get('fecha_cdp', '')}'")
    print(f"numero_cdp: '{row.get('numero_cdp', '')}'")
    print(f"cuenta_cdp: '{row.get('cuenta_cdp', '')}'")
    print(f"valor_cdp: '{row.get('valor_cdp', '')}'")
    
    doc = Document(TEMPLATE_PATH)

    # Asegurarse de que las claves coincidan con las columnas del Excel
    # Convertir nombres de columnas a minúsculas para evitar problemas de mayúsculas/minúsculas
    row_lower = {k.lower(): v for k, v in row.items()}
    
    reemplazos = {
        "modalidad": row_lower.get("modalidad", ""),
        "tipo_contratacion": row_lower.get("tipo_contratacion", ""),
        "numero": str(row_lower.get("numero", "")),
        "identificacion": str(row_lower.get("identificacion", "")),
        "razón_social": row_lower.get("razón_social", ""),
        "valor": formato_moneda(row_lower.get("valor", 0)),
        "fecha_suscripcion": row_lower.get("fecha_suscripcion", ""),
        "fecha_legalizacion": row_lower.get("fecha_legalizacion", ""),
        "fecha_cdp": row_lower.get("fecha_cdp", ""),
        "numero_cdp": str(row_lower.get("numero_cdp", "")),
        "cuenta_cdp": row_lower.get("cuenta_cdp", ""),
        "valor_cdp": formato_moneda(row_lower.get("valor_cdp", 0)),
        "fecha_rp": row_lower.get("fecha_rp", ""),
        "numero_rp": str(row_lower.get("numero_rp", "")),
        "cuenta_rp": row_lower.get("cuenta_rp", ""),
        "valor_rp": formato_moneda(row_lower.get("valor_rp", 0)),
        "NOMBRE_SUPERVISOR": row_lower.get("nombre_supervisor", ""),
        "CARGO_SUPERVISOR": row_lower.get("cargo_supervisor", "")
    }
    
    # Función mejorada para reemplazar texto en párrafos
    def replace_text_in_paragraph(paragraph, key, value):
        if f"«{key}»" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(f"«{key}»", str(value))
    
    # Reemplazar en todos los párrafos del documento
    for p in doc.paragraphs:
        for key, val in reemplazos.items():
            replace_text_in_paragraph(p, key, val)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in reemplazos.items():
                        replace_text_in_paragraph(p, key, val)

    # Guardar .docx temporal
    tmp_dir = tempfile.gettempdir()
    docx_path = os.path.join(tmp_dir, f"certificado_{row_id}.docx")
    pdf_path = docx_path.replace(".docx", ".pdf")
    doc.save(docx_path)

    # Convertir a PDF usando LibreOffice
    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", tmp_dir,
        docx_path
    ], check=True)

    return send_file(pdf_path, as_attachment=True, download_name=f"certificado_{row_id+1}.pdf")

if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    app.run(host="0.0.0.0", port=8080)
